"""
Document Service - 文档处理微服务
负责文档接收、预处理、解析和格式转换
端口: 8003
"""

import asyncio
import logging
import os
import hashlib
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
import json
import mimetypes

import aiofiles
import aiohttp
from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import asyncpg
import redis.asyncio as redis
from pydantic import BaseModel, Field

# 导入事件系统
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from microservices.event_system import EventBus, Event, EventType, create_event

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 环境变量配置
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://postgres:postgres@localhost:5432/knowledge_base")
REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
SERVICE_NAME = os.getenv("SERVICE_NAME", "document-service")
SERVICE_PORT = int(os.getenv("SERVICE_PORT", "8003"))
MAX_WORKERS = int(os.getenv("MAX_WORKERS", "4"))
DATA_DIR = os.getenv("DATA_DIR", "./data")
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")

# 确保数据目录存在
Path(DATA_DIR).mkdir(parents=True, exist_ok=True)


class DocumentUploadRequest(BaseModel):
    """文档上传请求模型"""
    filename: str
    file_type: str
    content_type: str
    file_size: int
    metadata: Optional[Dict[str, Any]] = {}


class DocumentProcessRequest(BaseModel):
    """文档处理请求模型"""
    document_id: str
    processing_options: Optional[Dict[str, Any]] = {}


class DocumentResponse(BaseModel):
    """文档响应模型"""
    id: str
    filename: str
    file_type: str
    status: str
    created_at: datetime
    updated_at: datetime
    file_hash: str
    file_size: int
    metadata: Dict[str, Any]
    processing_results: Optional[Dict[str, Any]] = None


class DocumentProcessor:
    """文档处理器"""

    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.ms-excel': self._process_xls,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'text/plain': self._process_text,
            'text/csv': self._process_csv,
            'image/jpeg': self._process_image,
            'image/png': self._process_image,
            'image/tiff': self._process_image
        }

    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """处理文档"""
        processor = self.supported_formats.get(file_type)
        if not processor:
            raise ValueError(f"Unsupported file type: {file_type}")

        try:
            result = await processor(file_path)
            return result
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise

    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """处理PDF文档"""
        try:
            import PyPDF2
            import fitz  # PyMuPDF

            # 使用PyMuPDF进行更好的PDF处理
            doc = fitz.open(file_path)
            text_content = ""
            images = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_content += page.get_text()

                # 提取图像
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha < 4:  # 确保是RGB或灰度图像
                        img_data = pix.tobytes("png")
                        images.append({
                            'page': page_num,
                            'index': img_index,
                            'data': img_data,
                            'size': len(img_data)
                        })
                    pix = None

            doc.close()

            # 使用PyPDF2作为备选
            if not text_content.strip():
                text_content = await self._extract_text_with_pypdf2(file_path)

            return {
                'text_content': text_content,
                'images': images,
                'page_count': doc.page_count if 'doc' in locals() else 0,
                'processing_method': 'PyMuPDF + PyPDF2',
                'confidence_score': 0.9 if text_content.strip() else 0.3
            }

        except ImportError:
            # 如果PyMuPDF不可用，使用PyPDF2
            return await self._extract_text_with_pypdf2(file_path)
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    async def _extract_text_with_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """使用PyPDF2提取文本"""
        import PyPDF2

        text_content = ""
        page_count = 0

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)

            for page in pdf_reader.pages:
                text_content += page.extract_text() or ""

        return {
            'text_content': text_content,
            'images': [],
            'page_count': page_count,
            'processing_method': 'PyPDF2',
            'confidence_score': 0.7 if text_content.strip() else 0.3
        }

    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """处理Word文档"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = ""
            tables = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                'text_content': text_content,
                'tables': tables,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'processing_method': 'python-docx',
                'confidence_score': 0.9
            }

        except ImportError:
            logger.error("python-docx not installed")
            raise
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    async def _process_xls(self, file_path: str) -> Dict[str, Any]:
        """处理Excel 97-2003文件"""
        try:
            import xlrd

            workbook = xlrd.open_workbook(file_path)
            sheets_data = []

            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                sheet_data = []

                for row_idx in range(sheet.nrows):
                    row_data = []
                    for col_idx in range(sheet.ncols):
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        row_data.append(str(cell_value) if cell_value else "")
                    sheet_data.append(row_data)

                sheets_data.append({
                    'name': sheet_name,
                    'data': sheet_data,
                    'rows': sheet.nrows,
                    'cols': sheet.ncols
                })

            return {
                'text_content': self._extract_text_from_excel_data(sheets_data),
                'sheets': sheets_data,
                'sheet_count': len(sheets_data),
                'processing_method': 'xlrd',
                'confidence_score': 0.9
            }

        except ImportError:
            logger.error("xlrd not installed")
            raise
        except Exception as e:
            logger.error(f"Error processing XLS {file_path}: {e}")
            raise

    async def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """处理Excel 2007+文件"""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, read_only=True)
            sheets_data = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []

                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(row_data)

                sheets_data.append({
                    'name': sheet_name,
                    'data': sheet_data,
                    'rows': sheet.max_row,
                    'cols': sheet.max_column
                })

            workbook.close()

            return {
                'text_content': self._extract_text_from_excel_data(sheets_data),
                'sheets': sheets_data,
                'sheet_count': len(sheets_data),
                'processing_method': 'openpyxl',
                'confidence_score': 0.9
            }

        except ImportError:
            logger.error("openpyxl not installed")
            raise
        except Exception as e:
            logger.error(f"Error processing XLSX {file_path}: {e}")
            raise

    def _extract_text_from_excel_data(self, sheets_data: List[Dict]) -> str:
        """从Excel数据中提取文本"""
        text_content = ""
        for sheet in sheets_data:
            text_content += f"Sheet: {sheet['name']}\n"
            for row_idx, row in enumerate(sheet['data']):
                row_text = "\t".join(row)
                text_content += f"Row {row_idx + 1}: {row_text}\n"
            text_content += "\n"
        return text_content

    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """处理纯文本文件"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()

            return {
                'text_content': content,
                'line_count': len(content.splitlines()),
                'character_count': len(content),
                'processing_method': 'aiofiles',
                'confidence_score': 1.0
            }

        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                async with aiofiles.open(file_path, 'r', encoding='gbk') as file:
                    content = await file.read()

                return {
                    'text_content': content,
                    'line_count': len(content.splitlines()),
                    'character_count': len(content),
                    'processing_method': 'aiofiles-gbk',
                    'confidence_score': 0.9
                }
            except Exception as e:
                logger.error(f"Failed to decode text file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise

    async def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """处理CSV文件"""
        try:
            import csv

            rows = []
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()

            reader = csv.reader(content.splitlines())
            for row in reader:
                rows.append(row)

            text_content = "\n".join([",".join(row) for row in rows])

            return {
                'text_content': text_content,
                'rows': rows,
                'row_count': len(rows),
                'col_count': len(rows[0]) if rows else 0,
                'processing_method': 'csv',
                'confidence_score': 0.9
            }

        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            raise

    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """处理图像文件"""
        try:
            from PIL import Image
            import pytesseract

            # 打开图像
            image = Image.open(file_path)

            # 提取文本
            try:
                text_content = pytesseract.image_to_string(image, lang='chi_sim+eng')
                confidence_score = 0.7 if text_content.strip() else 0.3
            except Exception as e:
                logger.warning(f"OCR failed for {file_path}: {e}")
                text_content = ""
                confidence_score = 0.1

            return {
                'text_content': text_content,
                'image_info': {
                    'format': image.format,
                    'mode': image.mode,
                    'size': image.size,
                    'file_size': os.path.getsize(file_path)
                },
                'processing_method': 'PIL + pytesseract',
                'confidence_score': confidence_score,
                'ocr_performed': True
            }

        except ImportError:
            logger.error("PIL or pytesseract not installed")
            return {
                'text_content': "",
                'image_info': {'format': 'unknown'},
                'processing_method': 'none',
                'confidence_score': 0.0,
                'ocr_performed': False,
                'error': 'Missing dependencies: PIL, pytesseract'
            }
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise


class DocumentService:
    """文档服务主类"""

    def __init__(self):
        self.app = FastAPI(
            title="Document Service",
            description="文档处理微服务",
            version="1.0.0"
        )
        self.db_pool = None
        self.redis_client = None
        self.event_bus = None
        self.processor = DocumentProcessor()

        # 配置CORS
        self.app.add_middleware(
            CORSMiddleware,
            allow_origins=["*"],
            allow_credentials=True,
            allow_methods=["*"],
            allow_headers=["*"],
        )

        # 注册路由
        self._register_routes()

    async def initialize(self):
        """初始化服务"""
        # 连接数据库
        self.db_pool = await asyncpg.create_pool(DATABASE_URL, min_size=2, max_size=10)
        logger.info("Connected to PostgreSQL database")

        # 连接Redis
        self.redis_client = redis.from_url(REDIS_URL, decode_responses=False)
        await self.redis_client.ping()
        logger.info("Connected to Redis")

        # 初始化事件总线
        self.event_bus = EventBus(REDIS_URL)
        await self.event_bus.connect()

        # 创建消费者组
        await self.event_bus.create_consumer_group(
            SERVICE_NAME,
            [EventType.DOCUMENT_UPLOADED, EventType.DOCUMENT_PROCESSING_FAILED]
        )

        # 启动事件消费者
        asyncio.create_task(
            self.event_bus.consume_events(
                SERVICE_NAME,
                [EventType.DOCUMENT_UPLOADED, EventType.DOCUMENT_PROCESSING_FAILED],
                self._handle_event
            )
        )

        logger.info("Document Service initialized successfully")

    def _register_routes(self):
        """注册API路由"""

        @self.app.on_event("startup")
        async def startup_event():
            await self.initialize()

        @self.app.on_event("shutdown")
        async def shutdown_event():
            if self.db_pool:
                await self.db_pool.close()
            if self.redis_client:
                await self.redis_client.close()
            if self.event_bus:
                await self.event_bus.disconnect()

        @self.app.get("/health")
        async def health_check():
            """健康检查"""
            return {
                "status": "healthy",
                "service": SERVICE_NAME,
                "timestamp": datetime.utcnow().isoformat(),
                "database": "connected" if self.db_pool else "disconnected",
                "redis": "connected" if self.redis_client else "disconnected"
            }

        @self.app.post("/upload", response_model=DocumentResponse)
        async def upload_document(
            file: UploadFile = File(...),
            background_tasks: BackgroundTasks = BackgroundTasks()
        ):
            """上传文档"""
            try:
                # 验证文件
                if not file.filename:
                    raise HTTPException(status_code=400, detail="No filename provided")

                # 计算文件哈希
                file_content = await file.read()
                file_hash = hashlib.md5(file_content).hexdigest()

                # 检查是否已存在
                existing_doc = await self._get_document_by_hash(file_hash)
                if existing_doc:
                    return DocumentResponse(**existing_doc)

                # 保存文件
                file_id = str(uuid.uuid4())
                file_extension = Path(file.filename).suffix
                saved_path = Path(DATA_DIR) / f"{file_id}{file_extension}"

                async with aiofiles.open(saved_path, 'wb') as f:
                    await f.write(file_content)

                # 创建文档记录
                document_data = {
                    'id': file_id,
                    'filename': file.filename,
                    'file_type': file.content_type,
                    'status': 'uploaded',
                    'created_at': datetime.utcnow(),
                    'updated_at': datetime.utcnow(),
                    'file_hash': file_hash,
                    'file_size': len(file_content),
                    'file_path': str(saved_path),
                    'metadata': {
                        'original_filename': file.filename,
                        'content_type': file.content_type,
                        'upload_source': 'api'
                    }
                }

                await self._save_document(document_data)

                # 发布文档上传事件
                event = create_event(
                    EventType.DOCUMENT_UPLOADED,
                    SERVICE_NAME,
                    {
                        'document_id': file_id,
                        'filename': file.filename,
                        'file_type': file.content_type,
                        'file_size': len(file_content),
                        'file_path': str(saved_path)
                    }
                )
                await self.event_bus.publish_event(event)

                # 启动后台处理
                background_tasks.add_task(self._process_document_background, file_id)

                return DocumentResponse(**document_data)

            except Exception as e:
                logger.error(f"Error uploading document: {e}")
                raise HTTPException(status_code=500, detail=str(e))

        @self.app.post("/process/{document_id}")
        async def process_document(document_id: str, request: DocumentProcessRequest):
            """处理文档"""
            try:
                # 获取文档信息
                document = await self._get_document(document_id)
                if not document:
                    raise HTTPException(status_code=404, detail="Document not found")

                if document['status'] == 'processing':
                    return {"status": "already_processing", "document_id": document_id}

                # 更新状态
                await self._update_document_status(document_id, 'processing')

                # 启动处理任务
                asyncio.create_task(
                    self._process_document_background(document_id, request.processing_options)
                )

                return {"status": "processing_started", "document_id": document_id}

            except Exception as e:
                logger.error(f"Error processing document {document_id}: {e}")
                await self._update_document_status(document_id, 'failed')
                raise HTTPException(status_code=500, detail=str(e))

        @self.app.get("/documents/{document_id}", response_model=DocumentResponse)
        async def get_document(document_id: str):
            """获取文档信息"""
            document = await self._get_document(document_id)
            if not document:
                raise HTTPException(status_code=404, detail="Document not found")

            return DocumentResponse(**document)

        @self.app.get("/documents")
        async def list_documents(
            skip: int = 0,
            limit: int = 50,
            status: Optional[str] = None,
            file_type: Optional[str] = None
        ):
            """列出文档"""
            documents = await self._list_documents(skip, limit, status, file_type)
            return {
                "documents": documents,
                "total": len(documents),
                "skip": skip,
                "limit": limit
            }

        @self.app.delete("/documents/{document_id}")
        async def delete_document(document_id: str):
            """删除文档"""
            try:
                document = await self._get_document(document_id)
                if not document:
                    raise HTTPException(status_code=404, detail="Document not found")

                # 删除文件
                file_path = document.get('file_path')
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)

                # 删除数据库记录
                await self._delete_document(document_id)

                return {"status": "deleted", "document_id": document_id}

            except Exception as e:
                logger.error(f"Error deleting document {document_id}: {e}")
                raise HTTPException(status_code=500, detail=str(e))

    async def _process_document_background(self, document_id: str, options: Optional[Dict] = None):
        """后台处理文档"""
        try:
            document = await self._get_document(document_id)
            if not document:
                logger.error(f"Document {document_id} not found")
                return

            await self._update_document_status(document_id, 'processing')

            file_path = document['file_path']
            file_type = document['file_type']

            # 处理文档
            processing_results = await self.processor.process_document(file_path, file_type)

            # 更新处理结果
            await self._update_processing_results(document_id, processing_results)
            await self._update_document_status(document_id, 'completed')

            # 发布处理完成事件
            event = create_event(
                EventType.DOCUMENT_PROCESSED,
                SERVICE_NAME,
                {
                    'document_id': document_id,
                    'filename': document['filename'],
                    'processing_results': processing_results,
                    'status': 'completed'
                }
            )
            await self.event_bus.publish_event(event)

            logger.info(f"Document {document_id} processed successfully")

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            await self._update_document_status(document_id, 'failed')

            # 发布处理失败事件
            event = create_event(
                EventType.DOCUMENT_PROCESSING_FAILED,
                SERVICE_NAME,
                {
                    'document_id': document_id,
                    'error': str(e),
                    'status': 'failed'
                }
            )
            await self.event_bus.publish_event(event)

    async def _handle_event(self, event: Event):
        """处理事件"""
        logger.info(f"Received event: {event.type.value} from {event.source}")

    # 数据库操作方法
    async def _save_document(self, document_data: Dict):
        """保存文档到数据库"""
        async with self.db_pool.acquire() as conn:
            await conn.execute(
                """
                INSERT INTO document_tasks (
                    id, file_path, file_hash, status, metadata, created_at
                ) VALUES ($1, $2, $3, $4, $5, $6)
                ON CONFLICT (id) DO UPDATE SET
                    status = EXCLUDED.status,
                    metadata = EXCLUDED.metadata,
                    updated_at = NOW()
                """,
                document_data['id'],
                document_data['file_path'],
                document_data['file_hash'],
                document_data['status'],
                json.dumps(document_data['metadata']),
                document_data['created_at']
            )

    async def _get_document(self, document_id: str) -> Optional[Dict]:
        """获取文档信息"""
        async with self.db_pool.acquire() as conn:
            row = await conn.fetchrow(
                "SELECT * FROM document_tasks WHERE id = $1",
                document_id
            )
            return dict(row) if row else None

    async def _get_document_by_hash(self, file_hash: str) -> Optional[Dict]:
        """根据哈希获取文档"""
        async with self.db_pool.acquire() as conn:
            row = await conn.fetchrow(
                "SELECT * FROM document_tasks WHERE file_hash = $1",
                file_hash
            )
            return dict(row) if row else None

    async def _update_document_status(self, document_id: str, status: str):
        """更新文档状态"""
        async with self.db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE document_tasks SET status = $1, updated_at = NOW() WHERE id = $2",
                status, document_id
            )

    async def _update_processing_results(self, document_id: str, results: Dict):
        """更新处理结果"""
        async with self.db_pool.acquire() as conn:
            await conn.execute(
                """
                UPDATE document_tasks
                SET metadata = jsonb_set(
                    metadata, '{processing_results}', $1
                ), updated_at = NOW()
                WHERE id = $2
                """,
                json.dumps(results), document_id
            )

    async def _list_documents(self, skip: int, limit: int, status: Optional[str], file_type: Optional[str]) -> List[Dict]:
        """列出文档"""
        query = "SELECT * FROM document_tasks WHERE 1=1"
        params = []
        param_count = 0

        if status:
            param_count += 1
            query += f" AND status = ${param_count}"
            params.append(status)

        if file_type:
            param_count += 1
            query += f" AND metadata->>'content_type' = ${param_count}"
            params.append(file_type)

        query += f" ORDER BY created_at DESC LIMIT ${param_count + 1} OFFSET ${param_count + 2}"
        params.extend([limit, skip])

        async with self.db_pool.acquire() as conn:
            rows = await conn.fetch(query, *params)
            return [dict(row) for row in rows]

    async def _delete_document(self, document_id: str):
        """删除文档"""
        async with self.db_pool.acquire() as conn:
            await conn.execute("DELETE FROM document_tasks WHERE id = $1", document_id)


# 创建应用实例
app_service = DocumentService()
app = app_service.app


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=SERVICE_PORT)